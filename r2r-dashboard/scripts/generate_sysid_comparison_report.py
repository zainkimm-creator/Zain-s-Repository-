"""Generate SysID plant/excitation comparison data, charts, and a DOCX report.

This script intentionally does not modify the dashboard UI. It uses the current
backend simulator and estimator so the generated report reflects the same model
behavior the dashboard exposes.
"""

from __future__ import annotations

import csv
import json
import math
import sys
from datetime import datetime
from pathlib import Path
from statistics import mean
from typing import Any, Iterable, Sequence

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from backend.models.simulation import SimulationConfig, simulate
from backend.sysid.estimator import estimate_parameters
from backend.validation.excitations import excitation_names, get_excitation_profile
from backend.validation.plants import parameters_for_plant, plant_registry


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = PROJECT_ROOT / "data" / "processed"
FIGURE_DIR = PROJECT_ROOT / "reports" / "figures"
SUMMARY_DIR = PROJECT_ROOT / "reports" / "validation_summary"
REPORT_PATH = SUMMARY_DIR / "sysid_plant_excitation_report.docx"
JSON_PATH = SUMMARY_DIR / "sysid_plant_excitation_report.json"
MATRIX_CSV = DATA_DIR / "sysid_plant_excitation_sweep.csv"
AMPLITUDE_CSV = DATA_DIR / "sysid_amplitude_sensitivity.csv"
TENSION_NOISE_CSV = DATA_DIR / "sysid_tension_noise_sensitivity.csv"
OMEGA_NOISE_CSV = DATA_DIR / "sysid_omega_noise_sensitivity.csv"
COMBINED_CSV = DATA_DIR / "sysid_report_all_cases.csv"

CHART_PATHS = {
    "heatmap": FIGURE_DIR / "sysid_plant_excitation_heatmap.png",
    "excitation_average": FIGURE_DIR / "sysid_excitation_average_bars.png",
    "best_by_plant": FIGURE_DIR / "sysid_best_excitation_by_plant.png",
    "amplitude": FIGURE_DIR / "sysid_amplitude_sensitivity.png",
    "tension_noise": FIGURE_DIR / "sysid_tension_noise_sensitivity.png",
    "omega_noise": FIGURE_DIR / "sysid_omega_noise_sensitivity.png",
}

BASELINE_COMPATIBLE_LIMIT_N = 12000.0
DEFAULT_DURATION_S = 4.0
DEFAULT_TLOG_MS = 10.0

TEXT = (27, 39, 55)
MUTED = (88, 99, 114)
GRID = (218, 225, 232)
BLUE = (41, 112, 219)
ORANGE = (229, 126, 36)
GREEN = (44, 151, 92)
PURPLE = (125, 91, 199)
RED = (201, 64, 64)


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    names = [
        "arialbd.ttf" if bold else "arial.ttf",
        "segoeuib.ttf" if bold else "segoeui.ttf",
    ]
    for name in names:
        path = Path("C:/Windows/Fonts") / name
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    left, top, right, bottom = draw.textbbox((0, 0), text, font=font)
    return right - left, bottom - top


def _draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    draw.text((42, 30), title, fill=TEXT, font=_font(30, True))
    if subtitle:
        draw.text((42, 70), subtitle, fill=MUTED, font=_font(18))


def _format_metric(value: Any, digits: int = 4) -> str:
    if value is None:
        return "n/a"
    if isinstance(value, str):
        return value
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(number):
        return "n/a"
    if abs(number) >= 1_000_000:
        return f"{number:.3e}"
    if 0 < abs(number) < 0.001:
        return f"{number:.3e}"
    if abs(number) >= 1000:
        return f"{number:,.0f}"
    if abs(number) >= 10:
        return f"{number:.2f}"
    return f"{number:.{digits}g}"


def _safe_float(value: Any) -> float | None:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    if not math.isfinite(number):
        return None
    return number


def _finite_values(rows: Iterable[dict[str, Any]], key: str) -> list[float]:
    return [value for row in rows if (value := _safe_float(row.get(key))) is not None]


def _csv_write(path: Path, rows: Sequence[dict[str, Any]]) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    fieldnames: list[str] = []
    for row in rows:
        for key in row:
            if key not in fieldnames:
                fieldnames.append(key)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)
    return str(path)


def _json_ready(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(key): _json_ready(item) for key, item in value.items()}
    if isinstance(value, list | tuple):
        return [_json_ready(item) for item in value]
    if isinstance(value, float):
        return value if math.isfinite(value) else None
    if isinstance(value, Path):
        return str(value)
    return value


def _relative_error_from_table(error_table: Sequence[dict[str, Any]], prefix: str) -> float | None:
    values = [
        abs(float(row["relative_error"]))
        for row in error_table
        if str(row["parameter"]).startswith(prefix)
    ]
    return mean(values) if values else None


def run_sysid_case(
    *,
    plant_id: str,
    excitation: str,
    amplitude_v: float,
    tension_noise_n: float,
    omega_noise_rad_s: float,
    duration_s: float = DEFAULT_DURATION_S,
    tlog_ms: float = DEFAULT_TLOG_MS,
    seed: int = 2026,
    scenario: str = "plant-excitation",
) -> dict[str, Any]:
    """Run one dashboard-model SysID case and return a flat result row."""

    params, plant = parameters_for_plant(plant_id)
    row: dict[str, Any] = {
        "scenario": scenario,
        "plant_id": plant["plant_id"],
        "material": plant["material"],
        "scale": plant["scale"],
        "regime": plant["regime"],
        "EA_N": float(plant["EA_N"]),
        "baseline_range_compatible": bool(plant["baseline_range_compatible"]),
        "recommended_excitation_amplitude_V": float(plant["recommended_excitation_amplitude_V"]),
        "excitation": excitation,
        "amplitude_V": float(amplitude_v),
        "tension_noise_N": float(tension_noise_n),
        "omega_noise_rad_s": float(omega_noise_rad_s),
        "duration_s": float(duration_s),
        "Tlog_ms": float(tlog_ms),
        "status": "ok",
        "input_status": "zero excitation" if abs(amplitude_v) < 1e-12 else "dynamic excitation",
    }
    try:
        sim = simulate(
            params,
            config=SimulationConfig(
                duration_s=duration_s,
                log_sample_time_s=tlog_ms / 1000.0,
                sensor_noise_tension_N=tension_noise_n,
                sensor_noise_omega_rad_s=omega_noise_rad_s,
                seed=seed,
                output_name="sysid_report_source.csv",
            ),
            excitation=get_excitation_profile(excitation, amplitude_v),
            write_output=False,
        )
        sysid = estimate_parameters(sim.rows, nominal_params=params, true_params=params, summary_name=None)
        row.update(
            {
                "RMSE_theta": sysid.rmse_theta,
                "samples": len(sim.rows),
                "tension_rmse_N": sim.metrics.get("tension_rmse_N"),
                "control_effort_rms_V": sim.metrics.get("control_effort_rms_V"),
                "EA_estimate_N": sysid.estimates.get("EA"),
                "EA_relative_error": next(
                    (
                        float(item["relative_error"])
                        for item in sysid.error_table
                        if item["parameter"] == "EA"
                    ),
                    None,
                ),
                "kt_mean_abs_relative_error": _relative_error_from_table(sysid.error_table, "kt_"),
                "kf_mean_abs_relative_error": _relative_error_from_table(sysid.error_table, "kf_"),
            }
        )
    except Exception as exc:  # noqa: BLE001 - report generation must keep failed cases.
        row.update(
            {
                "status": "error",
                "error": str(exc),
                "RMSE_theta": None,
                "samples": 0,
                "tension_rmse_N": None,
                "control_effort_rms_V": None,
                "EA_estimate_N": None,
                "EA_relative_error": None,
                "kt_mean_abs_relative_error": None,
                "kf_mean_abs_relative_error": None,
            }
        )
    return row


def run_matrix_sweep() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for plant in plant_registry():
        amplitude = float(plant["recommended_excitation_amplitude_V"])
        for excitation in excitation_names():
            rows.append(
                run_sysid_case(
                    plant_id=str(plant["plant_id"]),
                    excitation=excitation,
                    amplitude_v=amplitude,
                    tension_noise_n=0.0,
                    omega_noise_rad_s=0.0,
                    seed=2026,
                    scenario="plant-excitation",
                )
            )
    return rows


def run_sensitivity_sweeps() -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    amplitude_rows = [
        run_sysid_case(
            plant_id="P01",
            excitation="E_Toggle",
            amplitude_v=amplitude,
            tension_noise_n=0.0,
            omega_noise_rad_s=0.0,
            seed=3030,
            scenario="amplitude-sensitivity",
        )
        for amplitude in (0.0, 0.02, 0.04, 0.08, 0.12)
    ]
    tension_noise_rows = [
        run_sysid_case(
            plant_id="P01",
            excitation="E_Toggle",
            amplitude_v=0.08,
            tension_noise_n=noise,
            omega_noise_rad_s=0.0,
            seed=4040,
            scenario="tension-noise-sensitivity",
        )
        for noise in (0.0, 0.025, 0.05, 0.075, 0.10)
    ]
    omega_noise_rows = [
        run_sysid_case(
            plant_id="P01",
            excitation="E_Toggle",
            amplitude_v=0.08,
            tension_noise_n=0.0,
            omega_noise_rad_s=noise,
            seed=5050,
            scenario="omega-noise-sensitivity",
        )
        for noise in (0.0, 0.0025, 0.005, 0.008, 0.010)
    ]
    return amplitude_rows, tension_noise_rows, omega_noise_rows


def summarize_matrix(rows: Sequence[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    best_by_plant: list[dict[str, Any]] = []
    for plant in plant_registry():
        plant_rows = [row for row in rows if row["plant_id"] == plant["plant_id"]]
        valid = [row for row in plant_rows if _safe_float(row.get("RMSE_theta")) is not None]
        if not valid:
            best_by_plant.append(
                {
                    "plant_id": plant["plant_id"],
                    "material": plant["material"],
                    "scale": plant["scale"],
                    "EA_N": plant["EA_N"],
                    "regime": plant["regime"],
                    "recommended_amplitude_V": plant["recommended_excitation_amplitude_V"],
                    "best_excitation": "n/a",
                    "best_RMSE_theta": None,
                    "paper_interpretation": "No valid case.",
                }
            )
            continue
        best = min(valid, key=lambda row: float(row["RMSE_theta"]))
        zero_input = all(abs(float(row["amplitude_V"])) < 1e-12 for row in valid)
        best_by_plant.append(
            {
                "plant_id": plant["plant_id"],
                "material": plant["material"],
                "scale": plant["scale"],
                "EA_N": plant["EA_N"],
                "regime": plant["regime"],
                "recommended_amplitude_V": plant["recommended_excitation_amplitude_V"],
                "best_excitation": "zero-input policy" if zero_input else best["excitation"],
                "best_RMSE_theta": best["RMSE_theta"],
                "paper_interpretation": (
                    "Weak-excitation result; the supplement states exact per-roller arrays exist, but the dashboard reference currently applies only plant EA."
                    if zero_input
                    else "Dynamic SysID comparison against the simulator truth for this plant."
                ),
            }
        )

    averages: list[dict[str, Any]] = []
    for excitation in excitation_names():
        excitation_rows = [row for row in rows if row["excitation"] == excitation]
        all_values = _finite_values(excitation_rows, "RMSE_theta")
        compatible_rows = [
            row
            for row in excitation_rows
            if bool(row["baseline_range_compatible"])
            and abs(float(row["amplitude_V"])) > 1e-12
        ]
        compatible_values = _finite_values(compatible_rows, "RMSE_theta")
        averages.append(
            {
                "excitation": excitation,
                "all_plants_mean_RMSE_theta": mean(all_values) if all_values else None,
                "baseline_compatible_mean_RMSE_theta": mean(compatible_values)
                if compatible_values
                else None,
                "valid_cases": len(all_values),
                "dynamic_cases": len(compatible_values),
            }
        )
    return best_by_plant, averages


def summarize_sensitivity(rows: Sequence[dict[str, Any]], parameter_key: str, label: str) -> dict[str, Any]:
    valid = [row for row in rows if _safe_float(row.get("RMSE_theta")) is not None]
    if not valid:
        return {
            "factor": label,
            "best_setting": "n/a",
            "best_RMSE_theta": None,
            "worst_setting": "n/a",
            "worst_RMSE_theta": None,
            "interpretation": "No valid cases.",
        }
    best = min(valid, key=lambda row: float(row["RMSE_theta"]))
    worst = max(valid, key=lambda row: float(row["RMSE_theta"]))
    interpretation = {
        "amplitude_V": "Too little amplitude gives poor excitation; moderate amplitude improves the regression signal.",
        "tension_noise_N": "Tension noise corrupts finite-difference tension dynamics and EA estimation.",
        "omega_noise_rad_s": "Omega noise corrupts velocity derivatives and the kt/kf regression.",
    }[parameter_key]
    return {
        "factor": label,
        "best_setting": _format_metric(best[parameter_key]),
        "best_RMSE_theta": best["RMSE_theta"],
        "worst_setting": _format_metric(worst[parameter_key]),
        "worst_RMSE_theta": worst["RMSE_theta"],
        "interpretation": interpretation,
    }


def _color_interp(a: tuple[int, int, int], b: tuple[int, int, int], t: float) -> tuple[int, int, int]:
    t = max(0.0, min(1.0, t))
    return tuple(round(a[i] + (b[i] - a[i]) * t) for i in range(3))


def _heat_color(t: float) -> tuple[int, int, int]:
    if t < 0.5:
        return _color_interp((63, 171, 103), (255, 205, 86), t * 2.0)
    return _color_interp((255, 205, 86), (214, 73, 73), (t - 0.5) * 2.0)


def write_heatmap(rows: Sequence[dict[str, Any]], path: Path) -> str:
    plants = [plant["plant_id"] for plant in plant_registry()]
    excitations = list(excitation_names())
    width, height = 1270, 880
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    _draw_title(
        draw,
        "Plant x Excitation SysID Error Heatmap",
        "Cell values are RMSE_theta; lower is better. Color uses log scaling.",
    )
    left, top, cell_w, cell_h = 210, 140, 190, 58
    lookup = {(row["plant_id"], row["excitation"]): row for row in rows}
    values = _finite_values(rows, "RMSE_theta")
    log_values = [math.log10(value + 1e-5) for value in values]
    lo, hi = min(log_values), max(log_values)

    for col, excitation in enumerate(excitations):
        x = left + col * cell_w
        tw, _ = _text_size(draw, excitation, _font(17, True))
        draw.text((x + (cell_w - tw) / 2, top - 34), excitation, fill=TEXT, font=_font(17, True))

    for row_idx, plant_id in enumerate(plants):
        y = top + row_idx * cell_h
        plant = next(item for item in plant_registry() if item["plant_id"] == plant_id)
        draw.text((30, y + 10), f"{plant_id}  EA={_format_metric(plant['EA_N'], 3)}", fill=TEXT, font=_font(15, True))
        for col, excitation in enumerate(excitations):
            x = left + col * cell_w
            item = lookup[(plant_id, excitation)]
            value = _safe_float(item.get("RMSE_theta"))
            if value is None:
                color = (226, 232, 240)
                label = "error"
            else:
                score = 0.0 if hi == lo else (math.log10(value + 1e-5) - lo) / (hi - lo)
                color = _heat_color(score)
                label = _format_metric(value, 3)
            draw.rounded_rectangle([x, y, x + cell_w - 10, y + cell_h - 8], radius=8, fill=color, outline="white", width=2)
            label_font = _font(17, True)
            tw, th = _text_size(draw, label, label_font)
            draw.text((x + (cell_w - 10 - tw) / 2, y + (cell_h - 8 - th) / 2), label, fill=(17, 24, 39), font=label_font)

    legend_y = top + len(plants) * cell_h + 35
    draw.text((left, legend_y), "Lower error", fill=MUTED, font=_font(15))
    for i in range(160):
        color = _heat_color(i / 159.0)
        draw.line((left + 95 + i, legend_y + 8, left + 95 + i, legend_y + 28), fill=color)
    draw.text((left + 275, legend_y), "Higher error", fill=MUTED, font=_font(15))
    draw.text(
        (42, height - 58),
        "Note: P03-P10 use zero recommended excitation in the current dashboard model because their EA is outside the extracted baseline range.",
        fill=MUTED,
        font=_font(15),
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return str(path)


def _draw_axes(
    draw: ImageDraw.ImageDraw,
    *,
    left: int,
    top: int,
    right: int,
    bottom: int,
    ymax: float,
    ylabel: str,
) -> None:
    draw.line((left, bottom, right, bottom), fill=TEXT, width=2)
    draw.line((left, top, left, bottom), fill=TEXT, width=2)
    for i in range(5):
        y = bottom - (bottom - top) * i / 4
        value = ymax * i / 4
        draw.line((left, y, right, y), fill=GRID, width=1)
        draw.text((left - 82, y - 10), _format_metric(value, 3), fill=MUTED, font=_font(14))
    draw.text((left, top - 30), ylabel, fill=MUTED, font=_font(15))


def _draw_log_axes(
    draw: ImageDraw.ImageDraw,
    *,
    left: int,
    top: int,
    right: int,
    bottom: int,
    values: Sequence[float],
    ylabel: str,
) -> tuple[float, float]:
    positive = [value for value in values if value > 0]
    if not positive:
        positive = [1.0]
    log_min = math.floor(min(math.log10(value) for value in positive))
    log_max = math.ceil(max(math.log10(value) for value in positive))
    if log_max <= log_min:
        log_max = log_min + 1
    draw.line((left, bottom, right, bottom), fill=TEXT, width=2)
    draw.line((left, top, left, bottom), fill=TEXT, width=2)
    tick_count = min(7, log_max - log_min + 1)
    if tick_count <= 1:
        ticks = [log_min, log_max]
    else:
        step = (log_max - log_min) / (tick_count - 1)
        ticks = [log_min + step * idx for idx in range(tick_count)]
    for tick in ticks:
        y = bottom - (bottom - top) * (tick - log_min) / (log_max - log_min)
        draw.line((left, y, right, y), fill=GRID, width=1)
        draw.text((left - 92, y - 10), f"1e{int(round(tick))}", fill=MUTED, font=_font(14))
    draw.text((left, top - 30), ylabel, fill=MUTED, font=_font(15))
    return float(log_min), float(log_max)


def write_average_bar_chart(rows: Sequence[dict[str, Any]], path: Path) -> str:
    width, height = 1180, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    _draw_title(
        draw,
        "Average SysID Error by Excitation",
        "Grouped bars use a log scale because P02 creates large outlier errors.",
    )
    left, top, right, bottom = 105, 135, 1110, 580
    chart_values = _finite_values(rows, "all_plants_mean_RMSE_theta") + _finite_values(
        rows, "baseline_compatible_mean_RMSE_theta"
    )
    log_min, log_max = _draw_log_axes(
        draw,
        left=left,
        top=top,
        right=right,
        bottom=bottom,
        values=chart_values,
        ylabel="RMSE_theta (log scale)",
    )

    def y_to_px(value: float) -> float:
        log_value = math.log10(max(value, 1e-300))
        return bottom - (bottom - top) * (log_value - log_min) / (log_max - log_min)

    group_w = (right - left) / len(rows)
    bar_w = 42
    for idx, row in enumerate(rows):
        cx = left + group_w * idx + group_w / 2
        for offset, key, color, label in (
            (-bar_w / 2 - 4, "all_plants_mean_RMSE_theta", BLUE, "All plants"),
            (bar_w / 2 + 4, "baseline_compatible_mean_RMSE_theta", ORANGE, "Dynamic P01-P02"),
        ):
            value = _safe_float(row.get(key))
            if value is None:
                continue
            y_value = y_to_px(value)
            x0 = cx + offset - bar_w / 2
            x1 = cx + offset + bar_w / 2
            draw.rounded_rectangle([x0, y_value, x1, bottom], radius=6, fill=color)
            text = _format_metric(value, 3)
            tw, _ = _text_size(draw, text, _font(13, True))
            draw.text((x0 + (bar_w - tw) / 2, y_value - 23), text, fill=TEXT, font=_font(13, True))
        label = str(row["excitation"])
        tw, _ = _text_size(draw, label, _font(14, True))
        draw.text((cx - tw / 2, bottom + 16), label, fill=TEXT, font=_font(14, True))
    draw.rounded_rectangle([left + 15, 610, left + 32, 627], radius=3, fill=BLUE)
    draw.text((left + 42, 608), "All P01-P10", fill=MUTED, font=_font(14))
    draw.rounded_rectangle([left + 175, 610, left + 192, 627], radius=3, fill=ORANGE)
    draw.text((left + 202, 608), "Dynamic P01-P02 only", fill=MUTED, font=_font(14))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return str(path)


def write_best_by_plant_chart(rows: Sequence[dict[str, Any]], path: Path) -> str:
    width, height = 1240, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    _draw_title(
        draw,
        "Best Available SysID Result by Plant",
        "Bars use a log scale so P02 outlier behavior and low-error P01 behavior are both visible.",
    )
    left, top, right, bottom = 105, 135, 1165, 580
    log_min, log_max = _draw_log_axes(
        draw,
        left=left,
        top=top,
        right=right,
        bottom=bottom,
        values=_finite_values(rows, "best_RMSE_theta"),
        ylabel="RMSE_theta (log scale)",
    )

    def y_to_px(value: float) -> float:
        log_value = math.log10(max(value, 1e-300))
        return bottom - (bottom - top) * (log_value - log_min) / (log_max - log_min)

    group_w = (right - left) / len(rows)
    for idx, row in enumerate(rows):
        value = _safe_float(row.get("best_RMSE_theta"))
        if value is None:
            continue
        cx = left + group_w * idx + group_w / 2
        bar_w = min(54, group_w * 0.58)
        y_value = y_to_px(value)
        color = GREEN if row["recommended_amplitude_V"] else (151, 160, 171)
        draw.rounded_rectangle([cx - bar_w / 2, y_value, cx + bar_w / 2, bottom], radius=7, fill=color)
        text = _format_metric(value, 3)
        tw, _ = _text_size(draw, text, _font(13, True))
        draw.text((cx - tw / 2, y_value - 22), text, fill=TEXT, font=_font(13, True))
        plant_label = str(row["plant_id"])
        tw, _ = _text_size(draw, plant_label, _font(14, True))
        draw.text((cx - tw / 2, bottom + 14), plant_label, fill=TEXT, font=_font(14, True))
        excitation = str(row["best_excitation"])
        short = "zero" if "zero" in excitation else excitation
        tw, _ = _text_size(draw, short, _font(12))
        draw.text((cx - tw / 2, bottom + 36), short, fill=MUTED, font=_font(12))
    legend_y = 650
    draw.rounded_rectangle([left + 15, legend_y, left + 32, legend_y + 17], radius=3, fill=GREEN)
    draw.text((left + 42, legend_y - 2), "Dynamic excitation", fill=MUTED, font=_font(14))
    draw.rounded_rectangle([left + 185, legend_y, left + 202, legend_y + 17], radius=3, fill=(151, 160, 171))
    draw.text((left + 212, legend_y - 2), "Zero-input safety policy", fill=MUTED, font=_font(14))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return str(path)


def write_line_chart(
    rows: Sequence[dict[str, Any]],
    *,
    x_key: str,
    title: str,
    subtitle: str,
    x_label: str,
    path: Path,
    color: tuple[int, int, int],
) -> str:
    width, height = 1120, 680
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    _draw_title(draw, title, subtitle)
    left, top, right, bottom = 105, 132, 1045, 540
    valid = [
        row
        for row in rows
        if _safe_float(row.get(x_key)) is not None and _safe_float(row.get("RMSE_theta")) is not None
    ]
    x_values = [float(row[x_key]) for row in valid]
    y_values = [float(row["RMSE_theta"]) for row in valid]
    xmax, xmin = max(x_values), min(x_values)
    ymax = max(y_values) * 1.18 if y_values else 1.0
    _draw_axes(draw, left=left, top=top, right=right, bottom=bottom, ymax=ymax, ylabel="RMSE_theta")

    def x_to_px(x_value: float) -> float:
        if xmax == xmin:
            return (left + right) / 2
        return left + (right - left) * (x_value - xmin) / (xmax - xmin)

    def y_to_px(y_value: float) -> float:
        return bottom - (bottom - top) * y_value / ymax

    points = [(x_to_px(x), y_to_px(y)) for x, y in zip(x_values, y_values, strict=True)]
    if len(points) > 1:
        draw.line(points, fill=color, width=4)
    for (px, py), row in zip(points, valid, strict=True):
        draw.ellipse([px - 8, py - 8, px + 8, py + 8], fill=color, outline="white", width=2)
        text = _format_metric(row["RMSE_theta"], 3)
        tw, _ = _text_size(draw, text, _font(13, True))
        draw.text((px - tw / 2, py - 31), text, fill=TEXT, font=_font(13, True))
        x_text = _format_metric(row[x_key], 4)
        tw, _ = _text_size(draw, x_text, _font(13))
        draw.text((px - tw / 2, bottom + 16), x_text, fill=TEXT, font=_font(13))
    draw.text((left, bottom + 47), x_label, fill=MUTED, font=_font(15))
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return str(path)


def _set_cell_shading(cell: Any, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.6)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for style_name, size, color in (
        ("Title", 22, "17324D"),
        ("Heading 1", 15, "17324D"),
        ("Heading 2", 12, "2970DB"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _add_table(doc: Document, rows: Sequence[dict[str, Any]], columns: Sequence[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    for idx, (_, label) in enumerate(columns):
        cell = header.cells[idx]
        cell.text = label
        _set_cell_shading(cell, "E8F1FF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.size = Pt(8.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            cells[idx].text = _format_metric(row.get(key))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def _add_chart(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.9))
    caption_para = doc.add_paragraph(caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.runs[0].italic = True
    caption_para.runs[0].font.size = Pt(8.5)


def write_docx_report(summary: dict[str, Any]) -> str:
    doc = Document()
    _style_document(doc)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SysID Plant and Excitation Comparison Report")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Generated from the dashboard backend model on {summary['generated_at']}")
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.color.rgb = RGBColor(88, 99, 114)

    doc.add_heading("Table of Contents", level=1)
    for idx, item in enumerate(
        [
            "Objective and reference basis",
            "How the dashboard model generated the results",
            "Why SysID changes with plant, excitation, amplitude, and noise",
            "Plant and excitation comparison",
            "Sensitivity analysis",
            "Conclusion and stored artifacts",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{idx}. {item}")

    doc.add_heading("1. Objective and Reference Basis", level=1)
    doc.add_paragraph(
        "This report runs the current dashboard simulator and SysID estimator for plants P01-P10 across the "
        "available excitation profiles. It also runs simple sensitivity cases for excitation amplitude, tension "
        "sensor noise, and roller-speed noise."
    )
    doc.add_paragraph(
        "Reference basis: the supplementary material provides the ten plant EA values and damping/regime labels "
        "in Table S12, and the state-space matrix structure in Table S13. The extracted dashboard reference does "
        "not contain numeric per-plant arrays for R, J, f, L, or b, so the current dashboard model varies EA by "
        "plant and keeps the baseline arrays for those other parameters."
    )
    doc.add_paragraph(
        "Because the supplement does not provide a numeric SysID RMSE table for every plant/excitation pair, "
        "the comparison difference reported here is the estimator error against the simulator truth used by the "
        "dashboard model: RMSE_theta plus parameter relative-error columns in the stored CSV files."
    )

    doc.add_heading("2. How the Dashboard Model Generated the Results", level=1)
    doc.add_paragraph(
        "Each case runs a 4 s closed-loop simulation with 1 ms RK4 physics integration and 10 ms PLC logging. "
        "The SysID estimator then solves finite-difference regression equations for kt_UW, kt_Nip, kt_RW, "
        "kf_UW, kf_Nip, kf_RW, and EA."
    )
    doc.add_paragraph(
        "A lower RMSE_theta means the estimated parameter vector is closer to the simulator truth. For high-EA "
        "plants outside the extracted baseline range, the dashboard currently recommends zero excitation for "
        "numerical safety; those rows are useful weak-excitation checks, not full dynamic paper reproduction."
    )

    doc.add_heading("3. Why SysID Results Change", level=1)
    for text in summary["explanations"]:
        doc.add_paragraph(text, style=None)

    doc.add_heading("4. Plant and Excitation Comparison", level=1)
    _add_chart(
        doc,
        CHART_PATHS["heatmap"],
        "Figure 1. Heatmap of RMSE_theta for each plant/excitation pair. Lower values are better.",
    )
    _add_chart(
        doc,
        CHART_PATHS["excitation_average"],
        "Figure 2. Average RMSE_theta by excitation, shown for all plants and for the dynamic baseline-compatible cases.",
    )
    _add_chart(
        doc,
        CHART_PATHS["best_by_plant"],
        "Figure 3. Best available result per plant. Gray bars indicate zero-input safety-policy cases.",
    )
    doc.add_heading("Best Result by Plant", level=2)
    _add_table(
        doc,
        summary["best_by_plant"],
        [
            ("plant_id", "Plant"),
            ("material", "Material"),
            ("scale", "Scale"),
            ("EA_N", "EA (N)"),
            ("regime", "Regime"),
            ("recommended_amplitude_V", "Amp (V)"),
            ("best_excitation", "Best excitation"),
            ("best_RMSE_theta", "Best RMSE"),
        ],
    )
    doc.add_heading("Average Result by Excitation", level=2)
    _add_table(
        doc,
        summary["excitation_averages"],
        [
            ("excitation", "Excitation"),
            ("all_plants_mean_RMSE_theta", "All plants mean"),
            ("baseline_compatible_mean_RMSE_theta", "Dynamic P01-P02 mean"),
            ("valid_cases", "Valid cases"),
            ("dynamic_cases", "Dynamic cases"),
        ],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_heading("5. Sensitivity Analysis", level=1)
    _add_chart(
        doc,
        CHART_PATHS["amplitude"],
        "Figure 4. Amplitude sensitivity for P01 with E_Toggle.",
    )
    _add_chart(
        doc,
        CHART_PATHS["tension_noise"],
        "Figure 5. Tension-noise sensitivity for P01 with E_Toggle.",
    )
    _add_chart(
        doc,
        CHART_PATHS["omega_noise"],
        "Figure 6. Roller-speed-noise sensitivity for P01 with E_Toggle.",
    )
    _add_table(
        doc,
        summary["sensitivity_summary"],
        [
            ("factor", "Factor"),
            ("best_setting", "Best setting"),
            ("best_RMSE_theta", "Best RMSE"),
            ("worst_setting", "Worst setting"),
            ("worst_RMSE_theta", "Worst RMSE"),
            ("interpretation", "Interpretation"),
        ],
    )

    doc.add_heading("6. Conclusion and Stored Artifacts", level=1)
    for text in summary["conclusions"]:
        doc.add_paragraph(text)
    _add_table(
        doc,
        summary["artifacts"],
        [
            ("type", "Type"),
            ("description", "Description"),
            ("file", "File"),
        ],
    )

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return str(REPORT_PATH)


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    SUMMARY_DIR.mkdir(parents=True, exist_ok=True)

    matrix_rows = run_matrix_sweep()
    amplitude_rows, tension_noise_rows, omega_noise_rows = run_sensitivity_sweeps()
    best_by_plant, excitation_averages = summarize_matrix(matrix_rows)
    sensitivity_summary = [
        summarize_sensitivity(amplitude_rows, "amplitude_V", "Excitation amplitude"),
        summarize_sensitivity(tension_noise_rows, "tension_noise_N", "Tension noise"),
        summarize_sensitivity(omega_noise_rows, "omega_noise_rad_s", "Omega noise"),
    ]

    csv_paths = {
        "matrix": _csv_write(MATRIX_CSV, matrix_rows),
        "amplitude": _csv_write(AMPLITUDE_CSV, amplitude_rows),
        "tension_noise": _csv_write(TENSION_NOISE_CSV, tension_noise_rows),
        "omega_noise": _csv_write(OMEGA_NOISE_CSV, omega_noise_rows),
        "combined": _csv_write(COMBINED_CSV, matrix_rows + amplitude_rows + tension_noise_rows + omega_noise_rows),
    }

    chart_paths = {
        "heatmap": write_heatmap(matrix_rows, CHART_PATHS["heatmap"]),
        "excitation_average": write_average_bar_chart(excitation_averages, CHART_PATHS["excitation_average"]),
        "best_by_plant": write_best_by_plant_chart(best_by_plant, CHART_PATHS["best_by_plant"]),
        "amplitude": write_line_chart(
            amplitude_rows,
            x_key="amplitude_V",
            title="Amplitude Effect on SysID",
            subtitle="P01, E_Toggle, noise-free. Too little input gives weak parameter excitation.",
            x_label="Excitation amplitude (V)",
            path=CHART_PATHS["amplitude"],
            color=BLUE,
        ),
        "tension_noise": write_line_chart(
            tension_noise_rows,
            x_key="tension_noise_N",
            title="Tension Noise Effect on SysID",
            subtitle="P01, E_Toggle, amplitude 0.08 V. Noise affects tension dynamics and EA.",
            x_label="Tension sensor noise standard deviation (N)",
            path=CHART_PATHS["tension_noise"],
            color=ORANGE,
        ),
        "omega_noise": write_line_chart(
            omega_noise_rows,
            x_key="omega_noise_rad_s",
            title="Omega Noise Effect on SysID",
            subtitle="P01, E_Toggle, amplitude 0.08 V. Noise affects velocity derivatives.",
            x_label="Omega sensor noise standard deviation (rad/s)",
            path=CHART_PATHS["omega_noise"],
            color=PURPLE,
        ),
    }

    best_dynamic = min(
        (
            row
            for row in matrix_rows
            if bool(row["baseline_range_compatible"])
            and abs(float(row["amplitude_V"])) > 1e-12
            and _safe_float(row.get("RMSE_theta")) is not None
        ),
        key=lambda row: float(row["RMSE_theta"]),
    )
    best_all = min(
        (row for row in matrix_rows if _safe_float(row.get("RMSE_theta")) is not None),
        key=lambda row: float(row["RMSE_theta"]),
    )

    summary: dict[str, Any] = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "method": {
            "duration_s": DEFAULT_DURATION_S,
            "Tlog_ms": DEFAULT_TLOG_MS,
            "plants": len(plant_registry()),
            "excitations": list(excitation_names()),
            "matrix_cases": len(matrix_rows),
            "sensitivity_cases": len(amplitude_rows) + len(tension_noise_rows) + len(omega_noise_rows),
        },
        "best_dynamic_case": best_dynamic,
        "best_all_case": best_all,
        "best_by_plant": best_by_plant,
        "excitation_averages": excitation_averages,
        "sensitivity_summary": sensitivity_summary,
        "explanations": [
            "Excitation changes the result because SysID is a regression problem: different profiles excite different combinations of roller velocity and web-tension states. A profile that does not move the needed states makes the normal equations weak or ill-conditioned.",
            "Plant type changes the result because EA changes the stiffness and the speed-to-tension gain. The supplement also says each plant has different per-roller arrays, but the current dashboard reference only exposes EA numerically; R, J, f, L, and b remain baseline arrays in this run.",
            "Amplitude changes the signal-to-noise ratio. Too low amplitude gives weak parameter excitation, while too high amplitude can push the reduced model into biased or unstable behavior. The useful region is the moderate range where the state motion is visible but still controlled.",
            "Tension noise directly disturbs the finite-difference tension equation, so it mainly affects EA and any result that depends on tension derivatives.",
            "Omega noise disturbs roller-speed derivatives, so it mainly affects the kt and kf estimates from roller dynamics.",
        ],
        "conclusions": [
            f"The best dynamic baseline-compatible case was {best_dynamic['plant_id']} with {best_dynamic['excitation']} at RMSE_theta={_format_metric(best_dynamic['RMSE_theta'], 4)}.",
            "P01 and P02 are the only plants inside the extracted Table S4 baseline EA range, so they receive nonzero recommended excitation in the current dashboard model. P03-P10 are high-EA cases and are reported as zero-input safety-policy cases unless exact per-roller arrays and retuned excitation limits are supplied.",
            "The sensitivity plots show the practical SysID tradeoff: usable excitation amplitude improves identifiability, while tension noise and omega noise degrade estimates through the finite-difference equations.",
        ],
        "csv_paths": csv_paths,
        "chart_paths": chart_paths,
        "artifacts": [
            {"type": "CSV", "description": "Plant x excitation matrix", "file": str(MATRIX_CSV.name)},
            {"type": "CSV", "description": "Amplitude sensitivity", "file": str(AMPLITUDE_CSV.name)},
            {"type": "CSV", "description": "Tension-noise sensitivity", "file": str(TENSION_NOISE_CSV.name)},
            {"type": "CSV", "description": "Omega-noise sensitivity", "file": str(OMEGA_NOISE_CSV.name)},
            {"type": "CSV", "description": "All report cases combined", "file": str(COMBINED_CSV.name)},
            {"type": "PNG", "description": "Heatmap, bars, and sensitivity line charts", "file": "reports/figures/sysid_*.png"},
            {"type": "JSON", "description": "Machine-readable report summary", "file": str(JSON_PATH.name)},
            {"type": "DOCX", "description": "Word report", "file": str(REPORT_PATH.name)},
        ],
    }

    report_path = write_docx_report(summary)
    summary["report_path"] = report_path
    JSON_PATH.write_text(json.dumps(_json_ready(summary), indent=2), encoding="utf-8")

    print(json.dumps({"report_path": report_path, "json_path": str(JSON_PATH), **csv_paths, **chart_paths}, indent=2))


if __name__ == "__main__":
    main()
